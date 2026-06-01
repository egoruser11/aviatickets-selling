from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ТЗ_ChatGPT_таблицы_и_графы_РПЗ_система_продажи_авиабилетов.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "637083"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EEF7F6"
BORDER = "D9E0EA"
CODE_FILL = "F7F9FC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths)

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(39, 54, 75)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(9.5)

    set_table_geometry(table, widths)
    return table


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.add_run(text)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor(11, 95, 89)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    if "SubtitleCustom" not in doc.styles:
        subtitle = doc.styles.add_style("SubtitleCustom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = doc.styles["SubtitleCustom"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)


def build() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Техническое задание для ChatGPT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    doc.add_paragraph(
        "Отрисовка всех таблиц, графиков и схем РПЗ по Laravel-проекту «Система продажи авиабилетов»",
        style="SubtitleCustom",
    )

    add_caption(doc, "Таблица 1 - Паспорт задания")
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Проект", "Web-приложение «Система продажи авиабилетов»: каталог рейсов, покупка билетов пользователями, администрирование рейсов, билетов, пользователей и промокодов."],
            ["Стек", "Laravel, Blade, Laravel Sanctum, PostgreSQL, Docker Compose, web-сессии для интерфейса и API-токены Sanctum для внешних запросов."],
            ["Назначение файла", "Дать другому ChatGPT полный и недвусмысленный контекст, чтобы он отрисовал материалы для РПЗ: таблицы, графики, ER-схему, графы форм, статусы и отчетные показатели."],
            ["Роли", "Гость, покупатель user, администратор admin. Заблокированный покупатель остается авторизованным, но не может покупать, пополнять баланс и отменять билеты."],
            ["Основная БД", "PostgreSQL. SQLite допустим только как in-memory база для автоматических тестов, но не как целевая СУБД курсовой."],
            ["Авторизация", "Laravel Sanctum подключен: web-авторизация через session guard, API через personal access tokens. Swagger/OpenAPI не использовать как реализованный модуль."],
            ["Финансовая логика", "У пользователя есть внутренний учебный баланс. Покупка списывает деньги, отмена или удаление оплаченного рейса/билета возвращает фактически списанную сумму."],
            ["Банковская логика", "Бизнес-логика банковской системы, эквайринга, карт, платежных шлюзов и реальных банковских транзакций намеренно пропущена. Пополнение баланса считается учебной операцией внутри приложения."],
            ["Покупка мест", "Пользователь выбирает места из select-списков. За один присест можно купить до 6 мест. Места генерируются по схеме A-F на ряд и не вводятся вручную."],
        ],
        [2200, 7160],
    )

    add_callout(
        doc,
        "Ключевая установка",
        "Все таблицы, графы и схемы должны соответствовать реально реализованному Laravel-проекту: PostgreSQL, Sanctum, роли admin/user, внутренний учебный баланс, промокоды, несколько мест за покупку, возвраты денег и админская инструкция. Банковская система и эквайринг в проекте не реализуются. Не придумывать Swagger, оплату банковской картой, регистрацию рейсов через сторонние API или несуществующие сущности.",
    )

    doc.add_heading("1. Что должен получить ChatGPT на вход", level=1)
    doc.add_paragraph(
        "Ниже приведены вводные, которые нужно передать генератору диаграмм и таблиц. Они адаптированы под текущую реализацию проекта и формат РПЗ: минимум лишней теории, максимум проверяемых сущностей, правил и связей."
    )
    add_caption(doc, "Таблица 2 - Исходные вводные для генератора")
    add_table(
        doc,
        ["Блок", "Содержание"],
        [
            ["Публичная часть", "Главная страница, каталог будущих рейсов, карточка рейса, вход и регистрация покупателя."],
            ["Покупатель", "Регистрация, вход, просмотр рейсов, пополнение баланса, покупка одного или нескольких мест, ввод промокода, просмотр своих билетов, отмена билета с возвратом на баланс."],
            ["Администратор", "Панель администратора, справочник авиакомпаний, CRUD рейсов, CRUD билетов, управление пользователями, блокировка, корректировка баланса, промокоды, отчеты, страница-инструкция."],
            ["Баланс", "balance хранится в users. balance_transactions хранит подписанные суммы: положительные пополнения/возвраты, отрицательные покупки, админские корректировки."],
            ["Промокоды", "promo_codes содержит процентные и фиксированные скидки. На одну покупку применяется один код; он нормализуется к верхнему регистру, должен быть активен, попадать в период действия и не превышать max_uses. Конкретный аккаунт может использовать конкретный промокод только один раз; после использования код скрывается из списка предложений для этого аккаунта. Для админского CRUD действуют пределы: процент до 100%, фиксированная скидка до 100 000 руб., лимит до 10 000 применений."],
            ["Места", "seat_number выбирается из списка. Активно занятыми считаются только booked и paid. Cancelled сохраняется в истории, но место освобождает."],
            ["Отчеты", "Форма 4 фильтрует по авиакомпании, городу отправления, городу прибытия, статусу билета и датам; показывает количество билетов и оплаченную выручку."],
            ["Ограничения UI", "Ошибки валидации отображаются в Blade под полями и в общем блоке. Верхняя панель разделена на основное меню и админскую строку управления. HTML-ограничения форм дублируют серверные проверки, но серверная валидация остается обязательной."],
            ["Нормализация ввода", "Перед проверками удаляются случайные пробелы. Email приводятся к нижнему регистру, а коды промокодов, авиакомпаний и места билетов - к верхнему. Номер нового рейса пользователь не вводит: он автоматически формируется как код авиакомпании + ID рейса."],
        ],
        [2200, 7160],
    )

    doc.add_heading("2. Перечень обязательных таблиц для РПЗ", level=1)
    add_caption(doc, "Таблица 3 - Таблицы, которые нужно отрисовать в РПЗ")
    add_table(
        doc,
        ["ID", "Название", "Что раскрыть", "Требуемый формат"],
        [
            ["T1", "Матрица ролей и функций", "Гость, покупатель, администратор, заблокированный пользователь; доступные действия и ограничения.", "Markdown-таблица: роль, функция, маршрут/экран, ограничение."],
            ["T2", "Словарь данных БД", "Все таблицы и поля: тип, ключи, nullable, назначение.", "Разделить по users, airlines, flights, tickets, balance_transactions, promo_codes, promo_code_usages, personal_access_tokens."],
            ["T3", "Связи таблиц", "FK, кардинальности, правила cascade/nullOnDelete, бизнес-смысл связи.", "Таблица с отдельной строкой на каждую связь."],
            ["T4", "Бизнес-правила", "Баланс, места, промокоды, блокировка, возвраты, отчеты, валидация.", "Таблица: правило, где проверяется, что будет при нарушении."],
            ["T5", "Web-маршруты", "Публичные, пользовательские и административные Blade-маршруты.", "Группировка по ролям и middleware."],
            ["T6", "API-маршруты Sanctum", "Регистрация, вход, список рейсов, текущий пользователь, logout.", "Таблица method/path/controller/auth."],
            ["T7", "Экранные формы", "Главная, каталог, карточка рейса, баланс, билеты, админские CRUD, отчеты, промокоды, инструкция.", "Таблица: форма, роль, поля, действия, проверки."],
            ["T8", "Статусы и переходы", "Ticket: booked/paid/cancelled; promo usable/unusable; balance transaction types.", "Матрица или state-table."],
            ["T9", "Отчетные показатели", "Билеты, выручка по paid, фильтры, группировка по рейсу.", "Таблица показателей и формул."],
            ["T10", "Демо-данные", "admin@example.com, пользователи, три авиакомпании, три рейса, базовые промокоды.", "Краткая таблица для раздела тестирования."],
        ],
        [850, 1900, 4050, 2560],
    )

    doc.add_heading("3. Перечень графов, схем и графиков", level=1)
    add_caption(doc, "Таблица 4 - Графы, диаграммы и графики для РПЗ")
    add_table(
        doc,
        ["ID", "Название", "Нотация", "Содержание"],
        [
            ["G1", "ER-диаграмма", "Mermaid erDiagram", "users, airlines, flights, tickets, balance_transactions, promo_codes, promo_code_usages, personal_access_tokens; PK/FK и кардинальности."],
            ["G2", "Use Case", "Mermaid flowchart", "Актеры: гость, покупатель, администратор; сценарии: каталог, покупка, баланс, отмена, CRUD, блокировка, отчеты."],
            ["G3", "Покупка нескольких мест", "Mermaid flowchart TD", "Выбор рейса, выбор 1-6 мест, промокод, проверка баланса, создание tickets, списание balance, decrement seats_available."],
            ["G4", "Возврат средств", "Mermaid flowchart TD", "Отмена билета покупателем, удаление билета админом, удаление рейса админом; возврат price на баланс и запись refund."],
            ["G5", "Управление рейсом", "Mermaid sequenceDiagram", "Admin создает/редактирует/удаляет рейс; при удалении paid tickets возвращаются покупателям."],
            ["G6", "Состояния билета", "Mermaid stateDiagram-v2", "booked -> paid, paid -> cancelled, booked -> cancelled; cancelled не занимает место."],
            ["G7", "Жизненный цикл промокода", "Mermaid stateDiagram-v2", "active, inactive, not_started, expired, limit_reached; usable только при выполнении всех условий."],
            ["G8", "Баланс аккаунта", "Mermaid flowchart TD", "top_up, purchase, refund, admin adjustment; signed amount и запрет отрицательного баланса."],
            ["G9", "Граф экранных форм", "Mermaid flowchart LR", "Переходы guest/user/admin между страницами Blade."],
            ["G10", "Развертывание", "Mermaid flowchart", "Docker Compose: nginx -> php app -> postgres; Laravel app использует PostgreSQL."],
        ],
        [720, 2050, 1900, 4690],
    )

    doc.add_heading("4. Схема базы данных для включения в промпт", level=1)
    doc.add_paragraph(
        "В РПЗ рекомендуется показывать 7 бизнес-таблиц и отдельно служебную таблицу Sanctum. Это покрывает предметную область продажи авиабилетов, баланс аккаунта, промокоды, разовое использование промокода аккаунтом и авторизацию API без лишних несуществующих модулей."
    )

    dictionaries = {
        "users": [
            ["id", "bigint", "PK", "Нет", "Идентификатор пользователя."],
            ["name", "varchar", "-", "Нет", "Имя пользователя или администратора."],
            ["email", "varchar", "unique", "Нет", "Email для входа и связи."],
            ["email_verified_at", "timestamp", "-", "Да", "Метка подтверждения email; может не использоваться в курсовой."],
            ["password", "varchar", "-", "Нет", "Хеш пароля Laravel."],
            ["role", "varchar(16)", "index", "Нет", "Роль: user или admin. По умолчанию user."],
            ["balance", "decimal(10,2)", "-", "Нет", "Текущий баланс покупателя. Для admin обычно 0."],
            ["blocked_at", "timestamp", "index", "Да", "Если заполнено, пользователь заблокирован."],
            ["remember_token", "varchar", "-", "Да", "Служебный токен remember me."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты Laravel."],
        ],
        "airlines": [
            ["id", "bigint", "PK", "Нет", "Идентификатор авиакомпании."],
            ["name", "varchar", "-", "Нет", "Название авиакомпании."],
            ["code", "varchar(8)", "unique", "Нет", "Код авиакомпании: SU, S7, EK и т.п."],
            ["country", "varchar", "-", "Нет", "Страна регистрации авиакомпании."],
            ["phone", "varchar", "-", "Да", "Контактный телефон."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "flights": [
            ["id", "bigint", "PK", "Нет", "Идентификатор рейса."],
            ["airline_id", "bigint", "FK airlines.id", "Нет", "Авиакомпания; при удалении авиакомпании рейсы удаляются cascade."],
            ["flight_number", "varchar(16)", "unique", "Нет", "Автоматический номер рейса: код авиакомпании + '-' + ID рейса. Администратор не вводит его вручную."],
            ["origin", "varchar", "-", "Нет", "Город отправления."],
            ["destination", "varchar", "-", "Нет", "Город прибытия."],
            ["departure_at", "datetime", "-", "Нет", "Дата и время вылета."],
            ["arrival_at", "datetime", "-", "Нет", "Дата и время прилета; позже departure_at."],
            ["seats_total", "unsigned int", "-", "Нет", "Общее количество мест."],
            ["seats_available", "unsigned int", "-", "Нет", "Свободные места; 0 <= seats_available <= seats_total."],
            ["base_price", "decimal(10,2)", "-", "Нет", "Базовая цена одного билета до скидки."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "tickets": [
            ["id", "bigint", "PK", "Нет", "Идентификатор билета."],
            ["user_id", "bigint", "FK users.id", "Да", "Покупатель; при удалении пользователя становится NULL."],
            ["flight_id", "bigint", "FK flights.id", "Нет", "Рейс; при удалении рейса билеты удаляются cascade после возврата средств."],
            ["promo_code_id", "bigint", "FK promo_codes.id", "Да", "Промокод, примененный к покупке; при удалении промокода становится NULL."],
            ["passenger_name", "varchar", "-", "Нет", "Имя пассажира; в админском CRUD автоматически копируется из выбранного по email User."],
            ["passenger_email", "varchar", "-", "Нет", "Email пассажира; в админском CRUD автоматически копируется из выбранного User."],
            ["seat_number", "varchar(8)", "app-level unique", "Нет", "Место. Активная уникальность проверяется для booked/paid в рамках одного flight_id."],
            ["status", "varchar(16)", "index logically", "Нет", "booked, paid или cancelled."],
            ["price", "decimal(10,2)", "-", "Нет", "Фактически списанная сумма по этому билету после распределения скидки."],
            ["discount_amount", "decimal(10,2)", "-", "Нет", "Часть скидки, распределенная на этот билет."],
            ["purchased_at", "datetime", "-", "Нет", "Дата покупки/создания билета. При админском создании фиксируется сервером автоматически; при редактировании сохраняется."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "balance_transactions": [
            ["id", "bigint", "PK", "Нет", "Идентификатор операции."],
            ["user_id", "bigint", "FK users.id", "Нет", "Пользователь; при удалении пользователя операции удаляются cascade."],
            ["ticket_id", "bigint", "FK tickets.id", "Да", "Связанный билет; при удалении билета становится NULL."],
            ["type", "varchar(32)", "index", "Нет", "top_up, purchase, refund или admin."],
            ["amount", "decimal(10,2)", "-", "Нет", "Подписанная сумма: плюс для пополнений/возвратов, минус для покупки."],
            ["description", "varchar", "-", "Да", "Пояснение операции."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "promo_codes": [
            ["id", "bigint", "PK", "Нет", "Идентификатор промокода."],
            ["code", "varchar(32)", "unique", "Нет", "Код, вводимый покупателем; trim + uppercase, разрешены латинские буквы, цифры, дефис и нижнее подчеркивание."],
            ["name", "varchar", "-", "Да", "Описание промокода."],
            ["type", "varchar(16)", "-", "Нет", "percent или fixed."],
            ["value", "decimal(10,2)", "-", "Нет", "Процент 0,01-100% или фиксированная сумма 0,01-100 000 руб.; максимум зависит от type."],
            ["is_active", "boolean", "-", "Нет", "Признак включения промокода."],
            ["starts_at", "datetime", "-", "Да", "Дата начала действия."],
            ["expires_at", "datetime", "-", "Да", "Дата окончания действия."],
            ["max_uses", "unsigned int", "-", "Да", "Максимум использований: от 1 до 10 000 или NULL без общего лимита; при редактировании не меньше used_count."],
            ["used_count", "unsigned int", "-", "Нет", "Сколько раз промокод применили."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "promo_code_usages": [
            ["id", "bigint", "PK", "Нет", "Идентификатор факта использования промокода."],
            ["user_id", "bigint", "FK users.id", "Нет", "Аккаунт, который уже применил промокод; при удалении пользователя запись удаляется cascade."],
            ["promo_code_id", "bigint", "FK promo_codes.id", "Нет", "Использованный промокод; при удалении промокода запись удаляется cascade."],
            ["user_id + promo_code_id", "unique", "unique index", "Нет", "Гарантирует, что один аккаунт не сможет использовать один и тот же промокод повторно."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
        "personal_access_tokens": [
            ["id", "bigint", "PK", "Нет", "Идентификатор токена Sanctum."],
            ["tokenable_type, tokenable_id", "morphs", "index", "Нет", "Связь с моделью пользователя."],
            ["name", "text", "-", "Нет", "Название токена."],
            ["token", "varchar(64)", "unique", "Нет", "Хеш токена."],
            ["abilities", "text", "-", "Да", "Права токена."],
            ["last_used_at", "timestamp", "-", "Да", "Дата последнего использования."],
            ["expires_at", "timestamp", "index", "Да", "Дата истечения."],
            ["created_at, updated_at", "timestamp", "-", "Да", "Служебные даты."],
        ],
    }

    for table_name, rows in dictionaries.items():
        doc.add_heading(table_name, level=2)
        add_caption(doc, f"Таблица - Словарь данных `{table_name}`")
        add_table(doc, ["Поле", "Тип", "Ключ/индекс", "NULL", "Назначение"], rows, [1750, 1650, 1750, 850, 3360])

    doc.add_heading("5. Связи и ограничения", level=1)
    add_caption(doc, "Таблица 5 - Связи сущностей")
    add_table(
        doc,
        ["Связь", "Кардинальность", "Правило удаления", "Бизнес-смысл"],
        [
            ["airlines.id -> flights.airline_id", "1:N", "cascadeOnDelete", "Удаление авиакомпании удаляет ее рейсы."],
            ["flights.id -> tickets.flight_id", "1:N", "cascadeOnDelete", "Билет относится к одному рейсу; перед удалением рейса paid-билеты возвращаются на баланс."],
            ["users.id -> tickets.user_id", "1:N", "nullOnDelete", "Билет может сохранить историю пассажира даже если аккаунт удален."],
            ["promo_codes.id -> tickets.promo_code_id", "1:N", "nullOnDelete", "Билет хранит примененную скидку даже если промокод позже удалили."],
            ["users.id -> promo_code_usages.user_id", "1:N", "cascadeOnDelete", "Аккаунт хранит историю уже использованных промокодов."],
            ["promo_codes.id -> promo_code_usages.promo_code_id", "1:N", "cascadeOnDelete", "Промокод связан с фактами использования; unique(user_id, promo_code_id) запрещает повтор."],
            ["users.id -> balance_transactions.user_id", "1:N", "cascadeOnDelete", "История баланса принадлежит пользователю."],
            ["tickets.id -> balance_transactions.ticket_id", "1:N optional", "nullOnDelete", "Финансовая операция может ссылаться на билет, но не обязана."],
            ["users.id -> personal_access_tokens.tokenable", "1:N polymorphic", "cascade by model/token cleanup", "Sanctum API-токены принадлежат пользователю."],
        ],
        [2550, 1450, 1900, 3460],
    )

    add_caption(doc, "Таблица 6 - Бизнес-правила и инварианты")
    add_table(
        doc,
        ["Правило", "Где проверяется", "Что важно отразить в РПЗ"],
        [
            ["Покупать, пополнять баланс и отменять билеты может только auth user без blocked_at.", "middleware auth + not_blocked", "Заблокированный пользователь получает отказ/редирект и не создает финансовых операций."],
            ["Администратор имеет отдельный middleware admin и не покупает билеты через каталог.", "routes/web.php, EnsureAdmin", "Админ управляет справочниками и пользователями, покупатель покупает."],
            ["Баланс не может уйти в минус при покупке или админском списании.", "AccountController, Admin\\UserController", "Покупка отклоняется при недостатке средств; списание больше баланса запрещено."],
            ["За одну покупку выбирается 1-6 мест.", "AccountController validation", "Места приходят массивом seat_numbers, каждый элемент distinct."],
            ["Место выбирается только из списка сгенерированных мест рейса.", "Flight::seatNumbers(), Blade select", "Ручной ввод места в пользовательском сценарии запрещен."],
            ["Активная занятость места учитывает только booked и paid.", "TicketController, AccountController", "cancelled освобождает место, но сохраняется в истории."],
            ["Исторический cancelled-билет может хранить место, которое позже занял новый booked или paid-билет.", "TicketController validation", "Уникальность активного места проверяется только для booked и paid; история отмен не теряется."],
            ["Банковская подсистема не реализована.", "Архитектурное ограничение проекта", "Нет платежного шлюза, банковских счетов, карт, чеков, комиссий и статусов банковского платежа; баланс пополняется учебной операцией."],
            ["Промокод применим только если active, не раньше starts_at, не позже expires_at и не исчерпал max_uses.", "PromoCode::isUsable()", "used_count увеличивается один раз на успешную покупку, а не на каждый билет внутри покупки."],
            ["На одну покупку применяется один промокод.", "AccountController", "Комбинирование нескольких промокодов не предусмотрено."],
            ["Один аккаунт может использовать конкретный промокод только один раз.", "promo_code_usages unique(user_id, promo_code_id), AccountController", "Повторная попытка возвращает ошибку валидации поля promo_code."],
            ["Использованный аккаунтом промокод не показывается в списке предложений на карточке рейса.", "CatalogController::show", "Покупатель видит только доступные ему активные промокоды."],
            ["При создании и изменении промокода действуют разумные пределы.", "Admin\\PromoCodeController", "Процент: 0,01-100%; fixed: 0,01-100 000 руб.; max_uses: 1-10 000 или NULL. Нельзя уменьшить max_uses ниже used_count."],
            ["Код промокода нормализуется и проверяется до unique-запроса.", "Admin\\PromoCodeController", "Используются trim + uppercase; допустимы только A-Z, 0-9, дефис и нижнее подчеркивание. Повтор с другим регистром дает понятную ошибку формы, а не ошибку БД."],
            ["Строковые идентификаторы нормализуются перед валидацией.", "AuthController, Api\\AuthController, AirlineController, TicketController", "Email приводятся к lowercase; коды авиакомпаний и места - к uppercase; случайные пробелы удаляются."],
            ["Номер нового рейса создается автоматически.", "FlightController::store, FlightController::automaticFlightNumber", "После вставки рейса номер записывается как код авиакомпании + ID рейса, например SU-4. Ручной ввод и подмена номера запросом не используются."],
            ["При ручном добавлении билета админ выбирает email аккаунта.", "tickets/index.blade.php, TicketController", "Имя пассажира и user_id подтягиваются автоматически из выбранного User; клиентский запрос не может подменить имя или email."],
            ["Дата покупки билета фиксируется автоматически.", "TicketController", "При админском создании нового билета purchased_at устанавливается сервером через now(); в форме ручного ввода даты нет. При редактировании прежняя дата сохраняется."],
            ["Денежные поля принимают максимум два знака после запятой.", "AccountController, Admin\\UserController, FlightController, TicketController, Admin\\PromoCodeController", "Проверяются пополнение и корректировка баланса, цена рейса, цена билета и размер скидки."],
            ["Скидка не может быть больше общей суммы покупки.", "PromoCode::discountFor()", "Итоговая сумма = subtotal - min(discount, subtotal), после чего discount распределяется по созданным билетам."],
            ["Удаление рейса возвращает paid-билеты на баланс покупателей перед cascade delete.", "FlightController::destroy", "Это обязательная финансовая гарантия."],
            ["Удаление или админская отмена paid-билета возвращает price покупателю.", "TicketController", "Возврат фиксируется в balance_transactions type=refund."],
            ["Форма 4 использует select для origin/destination.", "ReportController + reports/index.blade.php", "Значения берутся из реальных flights."],
            ["Ошибки валидации показываются в Blade у конкретных полей.", "partials/field-error.blade.php", "В РПЗ отметить проверку удобства интерфейса."],
        ],
        [2500, 2400, 4460],
    )

    doc.add_heading("6. Маршруты и формы", level=1)
    add_caption(doc, "Таблица 7 - Web-маршруты")
    add_table(
        doc,
        ["Группа", "Метод и путь", "Имя", "Назначение"],
        [
            ["public", "GET /", "home", "Главная и сводка."],
            ["public", "GET /catalog", "catalog.index", "Каталог будущих рейсов."],
            ["public", "GET /catalog/{flight}", "catalog.show", "Карточка рейса и форма покупки."],
            ["guest", "GET/POST /login", "login, login.store", "Вход."],
            ["guest", "GET/POST /register", "register, register.store", "Регистрация покупателя."],
            ["auth", "POST /logout", "logout", "Выход."],
            ["auth+not_blocked", "GET/POST /account/balance", "account.balance, account.balance.top-up", "Баланс и пополнение."],
            ["auth+not_blocked", "GET /account/tickets", "account.tickets", "Мои билеты."],
            ["auth+not_blocked", "POST /catalog/{flight}/buy", "catalog.buy", "Покупка 1-6 мест с промокодом."],
            ["auth+not_blocked", "PATCH /account/tickets/{ticket}/cancel", "account.tickets.cancel", "Отмена билета и возврат."],
            ["auth+admin", "GET /admin", "admin.dashboard", "Панель администратора."],
            ["auth+admin", "GET /admin/help", "admin.help", "Инструкция администратора."],
            ["auth+admin", "GET/PATCH /admin/users...", "admin.users.*", "Пользователи, блокировка, баланс."],
            ["auth+admin", "resource /admin/promo-codes", "admin.promo-codes.*", "CRUD промокодов."],
            ["auth+admin", "resource /airlines", "airlines.*", "CRUD авиакомпаний."],
            ["auth+admin", "resource /flights", "flights.*", "CRUD рейсов с возвратом при удалении."],
            ["auth+admin", "GET /reports", "reports.index", "Форма 4: отчеты."],
            ["auth+admin", "resource /tickets", "tickets.*", "CRUD билетов."],
        ],
        [1500, 2600, 2200, 3060],
    )

    add_caption(doc, "Таблица 8 - API-маршруты Sanctum")
    add_table(
        doc,
        ["Метод", "Путь", "Контроллер", "Авторизация", "Назначение"],
        [
            ["POST", "/api/register", "Api\\AuthController@register", "Нет", "Регистрация API-пользователя и выдача токена."],
            ["POST", "/api/login", "Api\\AuthController@login", "Нет", "Вход и выдача персонального токена."],
            ["GET", "/api/flights", "Api\\FlightController@index", "Нет", "Публичный список будущих рейсов."],
            ["GET", "/api/user", "Closure", "auth:sanctum", "Текущий пользователь."],
            ["POST", "/api/logout", "Api\\AuthController@logout", "auth:sanctum", "Удаление текущего токена."],
        ],
        [900, 1800, 2700, 1500, 2460],
    )

    add_caption(doc, "Таблица 9 - Экранные формы")
    add_table(
        doc,
        ["Форма", "Роль", "Поля/данные", "Действия и проверки"],
        [
            ["Главная", "Все", "Сводка: авиакомпании, рейсы, билеты, покупатели.", "Переходы в каталог, аккаунт или админку."],
            ["Каталог рейсов", "Все", "Поиск по номеру, городу, авиакомпании; таблица рейсов.", "Открытие карточки рейса."],
            ["Карточка рейса", "Покупатель", "Пассажир, email, 6 select-мест, промокод, занятые места, предложения доступных промокодов.", "Проверка баланса, мест, промокода, скрытие уже использованных аккаунтом кодов, создание tickets."],
            ["Баланс", "Покупатель", "Сумма пополнения, история операций.", "top_up с min 100 max 500000."],
            ["Мои билеты", "Покупатель", "Рейс, маршрут, пассажир, место, статус, цена, скидка.", "Отмена active билета."],
            ["Пользователи", "Админ", "Поиск, статус, баланс, tickets_count.", "Блокировка, разблокировка, admin adjustment."],
            ["Рейсы CRUD", "Админ", "Авиакомпания, автоматический номер, маршрут, даты, seats_total, seats_available, base_price.", "Создание, изменение, удаление с возвратом paid-билетов. Номер формируется как код авиакомпании + ID рейса."],
            ["Билеты CRUD", "Админ", "Рейс, email аккаунта select, автоматически заполненное имя, select-место, статус, цена.", "Ручная запись/правка билета, контроль занятых мест. user_id, имя и email берутся из выбранного User; purchased_at нового билета фиксируется сервером автоматически."],
            ["Промокоды", "Админ", "Код, описание, тип percent/fixed, value, даты, лимит, active, used_count.", "CRUD; trim + uppercase кода; unique; percent до 100%; fixed до 100 000 руб.; max_uses до 10 000 и не ниже used_count."],
            ["Отчеты", "Админ", "Авиакомпания, origin select, destination select, статус, даты.", "Фильтр и вывод количества билетов/выручки."],
            ["Инструкция", "Админ", "Памятка по рейсам, билетам, пользователям, промокодам.", "Справочный экран без изменения данных."],
        ],
        [1700, 1100, 3300, 3260],
    )

    doc.add_heading("7. Статусы, переходы и финансовые операции", level=1)
    add_caption(doc, "Таблица 10 - Матрица статусов билета")
    add_table(
        doc,
        ["Статус", "Кто устанавливает", "Занимает место", "Финансовый смысл", "Допустимые переходы"],
        [
            ["booked", "Админ при ручном создании", "Да", "Бронь без обязательного списания баланса.", "booked -> paid, booked -> cancelled."],
            ["paid", "Покупатель при покупке; админ вручную", "Да", "Оплаченный билет. У покупателя списан price.", "paid -> cancelled."],
            ["cancelled", "Покупатель или админ", "Нет", "При переходе из paid создается refund на price.", "Конечный статус."],
        ],
        [1250, 2000, 1200, 2550, 2360],
    )

    add_caption(doc, "Таблица 11 - Типы операций баланса")
    add_table(
        doc,
        ["Тип", "Знак amount", "Когда создается", "Описание"],
        [
            ["top_up", "Плюс", "Покупатель пополняет баланс.", "Пополнение баланса пользователем."],
            ["purchase", "Минус", "Успешная покупка одного или нескольких мест.", "Списывается итоговая сумма с учетом промокода."],
            ["refund", "Плюс", "Отмена билета, удаление paid-билета или рейса.", "Возврат фактически списанной суммы price."],
            ["admin", "Плюс или минус", "Админ корректирует баланс.", "Списание не может превышать текущий баланс."],
        ],
        [1300, 1200, 3300, 3560],
    )

    add_caption(doc, "Таблица 12 - Промокоды")
    add_table(
        doc,
        ["Код из сида", "Тип", "Значение", "Правило"],
        [
            ["AVIA10", "percent", "10%", "Скидка 10% на общую сумму покупки."],
            ["STUDENT500", "fixed", "500 руб.", "Фиксированная скидка, но не больше общей суммы покупки."],
            ["FAMILY15", "percent", "15%", "Семейная скидка 15% на покупку нескольких билетов."],
            ["Любой admin code", "percent/fixed", "value из формы", "Должен быть active, не истекшим, не раньше starts_at, не исчерпавшим max_uses. Percent: до 100%; fixed: до 100 000 руб.; max_uses: до 10 000 и не ниже used_count. Код вводится один раз на всю покупку, не суммируется с другими кодами и может быть применен одним аккаунтом только один раз."],
        ],
        [1800, 1300, 1300, 4960],
    )

    doc.add_heading("8. Отчеты и демо-данные", level=1)
    add_caption(doc, "Таблица 13 - Отчетные показатели")
    add_table(
        doc,
        ["Показатель", "Формула/источник", "Фильтры"],
        [
            ["Количество билетов", "COUNT(tickets.id) по сгруппированному рейсу.", "airline_id, origin, destination, status, from_date, to_date."],
            ["Оплаченная выручка", "SUM(CASE WHEN tickets.status = 'paid' THEN tickets.price ELSE 0 END).", "Те же фильтры; booked и cancelled не дают выручку."],
            ["Рейсы", "flights join airlines left join tickets group by flights.id.", "Сортировка по departure_at."],
            ["Справочники фильтров", "airlines, distinct flights.origin, distinct flights.destination.", "Origin и destination выводятся select-списками."],
        ],
        [2200, 4300, 2860],
    )

    add_caption(doc, "Таблица 14 - Демо-данные и учетные записи")
    add_table(
        doc,
        ["Тип", "Данные", "Назначение"],
        [
            ["Администратор", "admin@example.com / password", "Проверка CRUD, отчетов, промокодов, блокировки."],
            ["Покупатель", "ivan.petrov@example.com / password", "Баланс 25000; есть paid-билет SU-100."],
            ["Покупатель", "anna.smirnova@example.com / password", "Баланс 8000; есть booked-билет S7-204."],
            ["Заблокированный", "oleg.volkov@example.com / password", "blocked_at заполнен; покупка и баланс недоступны."],
            ["Авиакомпании", "Аэрофлот, S7 Airlines, Emirates", "Справочник для рейсов."],
            ["Рейсы", "SU-100, S7-204, EK-132", "Три маршрута для каталога и отчетов."],
            ["Промокоды", "AVIA10, STUDENT500, FAMILY15", "Проверка percent/fixed скидок."],
        ],
        [1600, 3300, 4460],
    )

    doc.add_heading("9. Готовый промпт для ChatGPT", level=1)
    doc.add_paragraph("Этот блок можно копировать целиком. Он просит ChatGPT сгенерировать таблицы, Mermaid-графы и графики для вставки в РПЗ.")
    add_caption(doc, "Блок для копирования - основной промпт")
    add_code_block(
        doc,
        [
            "Ты выступаешь как системный аналитик и оформитель РПЗ для курсовой работы.",
            "Нужно отрисовать таблицы, графы, диаграммы и графики по web-приложению «Система продажи авиабилетов».",
            "",
            "Контекст проекта:",
            "- Стек: Laravel, Blade, Laravel Sanctum, PostgreSQL, Docker Compose.",
            "- Роли: guest, user, admin. Заблокированный user не может покупать билеты, пополнять баланс и отменять билеты.",
            "- User регистрируется, входит, смотрит каталог рейсов, выбирает 1-6 мест через select, вводит промокод, покупает билеты с баланса, смотрит свои билеты и отменяет их.",
            "- Admin управляет авиакомпаниями, рейсами, билетами, пользователями, балансами, промокодами, отчетами и страницей инструкции.",
            "- Swagger/OpenAPI не считать реализованным модулем. API есть только на Laravel Sanctum.",
            "- Банковская система намеренно не реализована: нет эквайринга, карт, платежных шлюзов и банковских транзакций. Пополнение баланса - учебная внутренняя операция приложения.",
            "",
            "Главные бизнес-правила:",
            "1. Целевая СУБД - PostgreSQL. SQLite используется только в автотестах.",
            "2. Пользовательский баланс хранится в users.balance, операции - в balance_transactions.",
            "3. Банковская логика пропущена: платежный шлюз, комиссии, банковские статусы и подтверждения карт не моделируются.",
            "4. Покупка создает по одному tickets на каждое выбранное место, списывает общую сумму и уменьшает flights.seats_available на количество мест.",
            "5. За одну покупку можно выбрать от 1 до 6 мест. Места должны быть distinct и входить в список Flight::seatNumbers() с буквами A-F.",
            "6. Активно занятые места - только tickets.status in (booked, paid). cancelled место освобождает.",
            "7. На одну покупку применяется не более одного промокода. Код нормализуется через trim + uppercase; допустимы A-Z, 0-9, дефис и нижнее подчеркивание.",
            "8. Промокод применим, если is_active=true, starts_at не в будущем, expires_at не в прошлом, used_count < max_uses или max_uses NULL.",
            "9. При админском создании и редактировании промокода percent ограничен диапазоном 0,01-100%, fixed - 0,01-100 000 руб., max_uses - 1-10 000 или NULL.",
            "10. При редактировании промокода нельзя установить max_uses меньше уже накопленного used_count.",
            "11. Один аккаунт может использовать конкретный промокод только один раз; это фиксируется в promo_code_usages unique(user_id, promo_code_id).",
            "12. Уже использованный аккаунтом промокод скрывается из списка предложений на карточке рейса.",
            "13. Скидка промокода не может превышать subtotal. Итог = subtotal - discount. discount распределяется по билетам.",
            "14. used_count увеличивается один раз на успешную покупку с промокодом.",
            "15. Денежные поля принимают не больше двух знаков после запятой. Email нормализуются в lowercase; коды авиакомпаний и места - в uppercase.",
            "16. Номер нового рейса не вводится вручную: FlightController формирует flight_number как код авиакомпании + '-' + ID рейса.",
            "17. При админском добавлении билета покупатель выбирается по email из select; user_id, passenger_name и passenger_email берутся из выбранного User и не доверяются клиентскому вводу.",
            "18. При админском создании билета purchased_at ставится сервером автоматически; ручного поля даты в форме нет. При редактировании существующего билета дата сохраняется.",
            "19. При недостаточном балансе покупка отклоняется; баланс не может уходить в минус.",
            "20. Отмена paid-билета возвращает ticket.price на баланс и создает balance_transactions type=refund.",
            "21. Удаление рейса администратором перед cascade delete возвращает paid-билеты покупателям.",
            "22. Удаление или админская отмена paid-билета тоже возвращает деньги.",
            "23. cancelled-билет сохраняется в истории и не блокирует место: новый booked или paid-билет может использовать тот же seat_number.",
            "24. Форма 4 использует select для авиакомпании, origin и destination и считает paid revenue только по tickets.status='paid'.",
            "",
            "Таблицы БД:",
            "users: id PK, name, email unique, email_verified_at nullable, password, role(user/admin) indexed, balance decimal(10,2), blocked_at nullable indexed, remember_token, timestamps.",
            "airlines: id PK, name, code unique varchar(8), country, phone nullable, timestamps.",
            "flights: id PK, airline_id FK airlines cascadeOnDelete, flight_number unique auto-generated as airline.code + '-' + flights.id, origin, destination, departure_at, arrival_at, seats_total, seats_available, base_price decimal(10,2), timestamps.",
            "tickets: id PK, user_id nullable FK users nullOnDelete, flight_id FK flights cascadeOnDelete, promo_code_id nullable FK promo_codes nullOnDelete, passenger_name, passenger_email, seat_number, status(booked/paid/cancelled), price decimal(10,2), discount_amount decimal(10,2), purchased_at, timestamps.",
            "balance_transactions: id PK, user_id FK users cascadeOnDelete, ticket_id nullable FK tickets nullOnDelete, type(top_up/purchase/refund/admin), amount signed decimal(10,2), description nullable, timestamps.",
            "promo_codes: id PK, code unique varchar(32), name nullable, type(percent/fixed), value decimal(10,2), is_active boolean, starts_at nullable, expires_at nullable, max_uses nullable, used_count, timestamps. В CRUD: code trim+uppercase и regex A-Z0-9_-, percent <= 100, fixed <= 100000, max_uses <= 10000 и не ниже used_count.",
            "promo_code_usages: id PK, user_id FK users cascadeOnDelete, promo_code_id FK promo_codes cascadeOnDelete, unique(user_id, promo_code_id), timestamps.",
            "personal_access_tokens: служебная таблица Sanctum: id PK, morphs tokenable, name, token unique, abilities nullable, last_used_at nullable, expires_at nullable indexed, timestamps.",
            "",
            "Нужно подготовить результат на русском языке:",
            "1. Markdown-таблицы: матрица ролей, словарь данных БД, связи, web/API маршруты, экранные формы, бизнес-правила, статусы, операции баланса, промокоды, отчетные показатели.",
            "2. Mermaid-диаграммы: ER, Use Case, покупка нескольких мест, возврат средств, управление рейсом, состояние билета, жизненный цикл промокода, баланс аккаунта, граф экранных форм, схема Docker-развертывания.",
            "3. Для каждой диаграммы дай короткую подпись, где объясняется, почему она нужна в РПЗ.",
            "4. Не добавляй сущности, которых нет в проекте: банковские платежи, банковские счета, карты, платежные шлюзы, аэропорты как отдельная таблица, Swagger, посадочные талоны, внешние API авиакомпаний.",
        ],
    )

    doc.add_heading("10. Контрольный список качества результата", level=1)
    add_bullets(
        doc,
        [
            "Все названия таблиц и полей должны совпадать с Laravel-миграциями.",
            "В ER-диаграмме обязательно показать nullable FK: tickets.user_id, tickets.promo_code_id, balance_transactions.ticket_id, а также таблицу promo_code_usages с unique(user_id, promo_code_id).",
            "В правилах мест указать, что cancelled не блокирует seat_number и может храниться в истории рядом с новым активным билетом на то же место.",
            "В финансовых правилах указать подписанный amount и возврат price, а не base_price.",
            "В финансовых правилах отдельно указать, что банковская система и эквайринг намеренно пропущены.",
            "В промокодах разделить percent и fixed, показать один код на покупку, max_uses, used_count, разовое использование на аккаунт, скрытие уже использованных кодов, распределение скидки по билетам и пределы CRUD: percent <= 100, fixed <= 100000, max_uses <= 10000.",
            "В маршрутах явно отделить public, guest, auth+not_blocked, auth+admin и auth:sanctum.",
            "В отчете выручки считать только paid tickets.",
            "В РПЗ не писать, что проект работает на SQLite; целевая СУБД - PostgreSQL.",
            "В UI-описании указать, что ошибки валидации отображаются под полями Blade, а HTML-ограничения форм дублируются серверной валидацией.",
            "В правилах ввода указать trim, lowercase для email, uppercase для кодов авиакомпаний, промокодов и мест.",
            "В описании CRUD рейсов указать, что flight_number автоматически формируется из кода авиакомпании и ID рейса; ручного ввода нет.",
            "В описании CRUD билетов указать email-select покупателя и серверное автозаполнение user_id, passenger_name и passenger_email из выбранного User.",
            "В описании CRUD билетов указать автоматическую серверную фиксацию purchased_at при создании и отсутствие ручного поля даты покупки.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
